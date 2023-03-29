from glob import glob
from tkinter import *
from tkinter import ttk
from tkinter import messagebox
from docx import Document



def replace_text_in_paragraph(paragraph, key, value):
            if key in paragraph.text:
                inline = paragraph.runs
                for item in inline:
                    if key in item.text:
                        item.text = item.text.replace(key, value)  


class DepTem3x20():
    def __init__(self):
        self.deptep3x20 = Tk()
        self.deptep3x20.geometry("800x800")
        self.deptep3x20.title("KepHelper")
        self.deptep3x20.protocol("WM_DELETE_WINDOW", self.on_closing)

        global FileName
        global CurrentDate
        global ClientCode
        global TreatyCode
        global AccTreatyPlacing
        global AccPlacingCurTag
        global ClientName
        global StateCode
        global DepositIBAN
        global DepositAccCurrency
        global DepositAmountInWords
        global DepositAmount
        global RateInWords
        global Rate
        global Err
        global Ett
        global Rtt
        global Dtt
        global Epp
        global Eww
        global Rqq
        global Rbb
        global Eqq
        global Ebb
        global DateFrom
        global DateInto
        global DayCount
        global PercentReturnType
        global AccTreatyPaymentPercent
        global AccPayBodyCurTag
        global ClientOpenDate

        Label(self.deptep3x20,text='FileName').place(x=250,y=15)
        FileName = Entry(self.deptep3x20,width=50)
        FileName.place(x=250, y=35) 
        Label(self.deptep3x20, text="CurrentDate").place(x=50,y=60)
        CurrentDate = Entry(self.deptep3x20)
        CurrentDate.place(x=50, y=80)
        Label(self.deptep3x20, text="ClientCode").place(x=50,y=120)
        ClientCode = Entry(self.deptep3x20)
        ClientCode.place(x=50, y=140)
        Label(self.deptep3x20, text="TreatyCode").place(x=50,y=180)
        TreatyCode = Entry(self.deptep3x20)
        TreatyCode.place(x=50, y=200)
        Label(self.deptep3x20, text="AccTreatyPlacing").place(x=50,y=240)
        AccTreatyPlacing = Entry(self.deptep3x20)
        AccTreatyPlacing.place(x=50, y=260)
        Label(self.deptep3x20, text="AccPlacingCurTag").place(x=50,y=300)
        AccPlacingCurTag = Entry(self.deptep3x20)
        AccPlacingCurTag.place(x=50, y=320)
        Label(self.deptep3x20, text="ClientName").place(x=50,y=360)
        ClientName = Entry(self.deptep3x20)
        ClientName.place(x=50, y=380)
        Label(self.deptep3x20, text="StateCode").place(x=50,y=420)
        StateCode = Entry(self.deptep3x20)
        StateCode.place(x=50, y=440)
        Label(self.deptep3x20, text="DepositIBAN").place(x=50,y=480)
        DepositIBAN = Entry(self.deptep3x20)
        DepositIBAN.place(x=50, y=500)
        Label(self.deptep3x20, text="DepositAccCurrency").place(x=50,y=540)
        DepositAccCurrency = Entry(self.deptep3x20)
        DepositAccCurrency.place(x=50, y=560)
        Label(self.deptep3x20, text="DepositAmount").place(x=50,y=600)
        DepositAmount = Entry(self.deptep3x20)
        DepositAmount.place(x=50, y=620)
        Label(self.deptep3x20, text="DepositAmountInWords").place(x=250,y=60)
        DepositAmountInWords = Entry(self.deptep3x20)
        DepositAmountInWords.place(x=250, y=80)
        Label(self.deptep3x20, text="Rate1").place(x=250,y=120)
        Rate = Entry(self.deptep3x20)
        Rate.place(x=250, y=140)
        Label(self.deptep3x20, text="Rate1InWords").place(x=250,y=180)
        RateInWords = Entry(self.deptep3x20)
        RateInWords.place(x=250, y=200)
        Label(self.deptep3x20, text="DateFrom1").place(x=250,y=240)
        Err = Entry(self.deptep3x20)
        Err.place(x=250, y=260)
        Label(self.deptep3x20, text="DateInto1").place(x=250,y=300)
        Ett = Entry(self.deptep3x20)
        Ett.place(x=250, y=320)
        Label(self.deptep3x20, text="Rate2").place(x=250,y=360)
        Rtt = Entry(self.deptep3x20)
        Rtt.place(x=250, y=380)
        Label(self.deptep3x20, text="Rate2InWords").place(x=250,y=420)
        Dtt = Entry(self.deptep3x20)
        Dtt.place(x=250, y=440)
        Label(self.deptep3x20, text="DateFrom2").place(x=250,y=480)
        Epp = Entry(self.deptep3x20)
        Epp.place(x=250, y=500)
        Label(self.deptep3x20, text="DateInto2").place(x=250,y=540)
        Eww = Entry(self.deptep3x20)
        Eww.place(x=250, y=560)
        Label(self.deptep3x20, text="Rate3").place(x=250,y=600)
        Rqq = Entry(self.deptep3x20)
        Rqq.place(x=250, y=620)
        Label(self.deptep3x20, text="Rate3InWords").place(x=500,y=60)
        Rbb = Entry(self.deptep3x20)
        Rbb.place(x=500, y=80)
        Label(self.deptep3x20, text="DateFrom3").place(x=500,y=120)
        Eqq = Entry(self.deptep3x20)
        Eqq.place(x=500, y=140)
        Label(self.deptep3x20, text="DateInto3").place(x=500,y=180)
        Ebb = Entry(self.deptep3x20)
        Ebb.place(x=500, y=200)
        Label(self.deptep3x20, text="DateFrom").place(x=500,y=240)
        DateFrom = Entry(self.deptep3x20)
        DateFrom.place(x=500, y=260)
        Label(self.deptep3x20, text="DateInto").place(x=500,y=300)
        DateInto = Entry(self.deptep3x20)
        DateInto.place(x=500, y=320)
        Label(self.deptep3x20, text="DayCount").place(x=500,y=360)
        DayCount = Entry(self.deptep3x20)
        DayCount.place(x=500, y=380)
        Label(self.deptep3x20, text="PercentReturnType").place(x=500,y=420)
        PercentReturnType = Entry(self.deptep3x20)
        PercentReturnType.place(x=500, y=440)
        Label(self.deptep3x20, text="AccTreatyPaymentPercent").place(x=500,y=480)
        AccTreatyPaymentPercent = Entry(self.deptep3x20)
        AccTreatyPaymentPercent.place(x=500, y=500)
        Label(self.deptep3x20, text="AccPayBodyCurTag").place(x=500,y=540)
        AccPayBodyCurTag = Entry(self.deptep3x20)
        AccPayBodyCurTag.place(x=500, y=560)
        Label(self.deptep3x20, text="ClientOpenDate").place(x=500,y=600)
        ClientOpenDate = Entry(self.deptep3x20)
        ClientOpenDate.place(x=500, y=620)
        Button(self.deptep3x20,text="Replace Text",command=self.search,height=3,width=10).place(x= 500,y=730)



    def search(self):
        template_file_path = 'C:\Test\DepositTemplate3x20.docx'
        output_file_path = 'C:\Test\\' + FileName.get() + '.docx'

        variables = {
        "{":"",
        "}":"",
        "CurrentDate": CurrentDate.get(),
        "ClientCode": ClientCode.get(),
        "TreatyCode": TreatyCode.get(),
        "AccTreatyPlacing": AccTreatyPlacing.get(),
        "AccPlacingCurTag": AccPlacingCurTag.get(),
        "ClientName": ClientName.get(),
        "StateCode": StateCode.get(),
        "DepositIBAN": DepositIBAN.get(),
        "DepositAccCurrency": DepositAccCurrency.get(),
        "DepositAmountInWords": DepositAmountInWords.get(),
        "DepositAmount": DepositAmount.get(),
        "RateInWords": RateInWords.get(),
        "Rate": Rate.get(),
        "Err": Err.get(),
        "Ett": Ett.get(),
        "Rtt": Rtt.get(),
        "Dtt": Dtt.get(),
        "Epp": Epp.get(),
        "Eww": Eww.get(),
        "Rqq": Rqq.get(),
        "Rbb":Rbb.get(),
        "Eqq":Eqq.get(),
        "Ebb":Ebb.get(),
        "DateFrom":DateFrom.get(),
        "DateInto":DateInto.get(),
        "DayCount":DayCount.get(),
        "PercentReturnType": PercentReturnType.get(),
        "AccTreatyPaymentPercent": AccTreatyPaymentPercent.get(),
        "AccPayBodyCurTag": AccPayBodyCurTag.get(),
        "ClientOpenDate":ClientOpenDate.get()
        }
        template_document = Document(template_file_path)

        for variable_key, variable_value in variables.items():
            for paragraph in template_document.paragraphs:
                replace_text_in_paragraph(paragraph, variable_key, variable_value)

            for table in template_document.tables:
                for col in table.columns:
                    for cell in col.cells:
                        for paragraph in cell.paragraphs:
                            replace_text_in_paragraph(paragraph, variable_key, variable_value)

        template_document.save(output_file_path)
        msg_box = messagebox.showinfo("Успішно","Заміна у файлі успішно виконанна")
        if msg_box == 'ok':
            self.deptep3x20.destroy()
    def on_closing(selt):
        if messagebox.askokcancel("Вийти?","Ви бажаєте закрити додаток?"):
            selt.deptep3x20.destroy()
            
            
    
    def start(self):
           self.deptep3x20.mainloop()
