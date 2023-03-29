from tkinter import *
from tkinter import ttk
from tkinter import messagebox
from docx import Document





def replace_text_in_paragraph(paragraph, key, value):
            if key in paragraph.text:
                inline = paragraph.runs
                for item in inline:
                    if key in item.text:
                        item.text = item.text.replace(key, value)  
class DepTU():
    def __init__(self):
        self.deptempusual = Tk()
        self.deptempusual.geometry("800x800")
        self.deptempusual.title("KepHelper")
        self.deptempusual.protocol("WM_DELETE_WINDOW", self.on_closing)


        global FileName
        global CurrentDate
        global ClientCode
        global TreatyCode
        global ClientName
        global StateCode
        global ClientAddress
        global AccTreatyPlacing
        global AccPlacingCurTag
        global DepositIBAN
        global DepositAccCurrency
        global DepositAmount
        global DepositAmountInWords
        global Rate
        global RateInWords
        global DateFrom
        global DateInto
        global DayCount
        global AutolongTreaty
        global PercentReturnType
        global AccTreatyPaymentPercent
        global AccPayBodyCurTag

        Label(self.deptempusual,text='FileName').place(x=250,y=15)
        FileName = Entry(self.deptempusual,width=50)
        FileName.place(x=250, y=35)
        Label(self.deptempusual, text="CurrentDate").place(x=50,y=60)
        CurrentDate = Entry(self.deptempusual)
        CurrentDate.place(x=50, y=80)
        Label(self.deptempusual, text="ClientCode").place(x=50,y=120)
        ClientCode = Entry(self.deptempusual)
        ClientCode.place(x=50, y=140)
        Label(self.deptempusual, text="TreatyCode").place(x=50,y=180)
        TreatyCode = Entry(self.deptempusual)
        TreatyCode.place(x=50, y=200)
        Label(self.deptempusual, text="ClientName").place(x=50,y=240)
        ClientName = Entry(self.deptempusual)
        ClientName.place(x=50, y=260)
        Label(self.deptempusual, text="StateCode").place(x=50,y=300)
        StateCode = Entry(self.deptempusual)
        StateCode.place(x=50, y=320)
        Label(self.deptempusual, text="ClientAddress").place(x=50,y=360)
        ClientAddress = Entry(self.deptempusual)
        ClientAddress.place(x=50, y=380)
        Label(self.deptempusual, text="AccTreatyPlacing").place(x=50,y=420)
        AccTreatyPlacing = Entry(self.deptempusual)
        AccTreatyPlacing.place(x=50, y=440)
        Label(self.deptempusual, text="AccPlacingCurTag").place(x=50,y=480)
        AccPlacingCurTag = Entry(self.deptempusual)
        AccPlacingCurTag.place(x=50, y=500)
        Label(self.deptempusual, text="DepositIBAN").place(x=50,y=540)
        DepositIBAN = Entry(self.deptempusual)
        DepositIBAN.place(x=50, y=560)
        Label(self.deptempusual, text="DepositAccCurrency").place(x=50,y=600)
        DepositAccCurrency = Entry(self.deptempusual)
        DepositAccCurrency.place(x=50, y=620)
        Label(self.deptempusual, text="DepositAmount").place(x=250,y=60)
        DepositAmount = Entry(self.deptempusual)
        DepositAmount.place(x=250, y=80)
        Label(self.deptempusual, text="DepositAmountInWords").place(x=250,y=120)
        DepositAmountInWords = Entry(self.deptempusual)
        DepositAmountInWords.place(x=250, y=140)
        Label(self.deptempusual, text="Rate").place(x=250,y=180)
        Rate = Entry(self.deptempusual)
        Rate.place(x=250, y=200)
        Label(self.deptempusual, text="RateInWords").place(x=250,y=240)
        RateInWords = Entry(self.deptempusual)
        RateInWords.place(x=250, y=260)
        Label(self.deptempusual, text="DateFrom").place(x=250,y=300)
        DateFrom = Entry(self.deptempusual)
        DateFrom.place(x=250, y=320)
        Label(self.deptempusual, text="DateInto").place(x=250,y=360)
        DateInto = Entry(self.deptempusual)
        DateInto.place(x=250, y=380)
        Label(self.deptempusual, text="DayCount").place(x=250,y=420)
        DayCount = Entry(self.deptempusual)
        DayCount.place(x=250, y=440)
        Label(self.deptempusual, text="AutolongTreaty").place(x=250,y=480)
        AutolongTreaty = Entry(self.deptempusual)
        AutolongTreaty.place(x=250, y=500)
        Label(self.deptempusual, text="PercentReturnType").place(x=250,y=540)
        PercentReturnType = Entry(self.deptempusual)
        PercentReturnType.place(x=250, y=560)
        Label(self.deptempusual, text="AccTreatyPaymentPercent").place(x=250,y=600)
        AccTreatyPaymentPercent = Entry(self.deptempusual)
        AccTreatyPaymentPercent.place(x=250, y=620)
        Label(self.deptempusual, text="AccPayBodyCurTag").place(x=500,y=60)
        AccPayBodyCurTag = Entry(self.deptempusual)
        AccPayBodyCurTag.place(x=500, y=80)
        Button(self.deptempusual,text="Replace Text",command = self.search,height=3,width=10).place(x= 500,y=730)
    
    def search(self):
        template_file_path = 'C:\Test\DepositTemplateUsual.docx'
        output_file_path = 'C:\Test\\' + FileName.get() + '.docx'

        variables = {
        "{":"",
        "}":"",
        "CurrentDate": CurrentDate.get(),
        "ClientCode": ClientCode.get(),
        "TreatyCode": TreatyCode.get(),
        "ClientName": ClientName.get(),
        "StateCode": StateCode.get(),
        "ClientAddress": ClientAddress.get(),
        "AccTreatyPlacing": AccTreatyPlacing.get(),
        "AccPlacingCurTag":AccPlacingCurTag.get(),
        "DepositIBAN":DepositIBAN.get(),
        "DepositAccCurrency":DepositAccCurrency.get(),
        "DepositAmountInWords":DepositAmountInWords.get(),
        "DepositAmount":DepositAmount.get(),
        "RateInWords":RateInWords.get(),
        "Rate":Rate.get(),
        "DateFrom":DateFrom.get(),
        "DateInto":DateInto.get(),
        "DayCount":DayCount.get(),
        "AutolongTreaty":AutolongTreaty.get(),
        "PercentReturnType":PercentReturnType.get(),
        "AccTreatyPaymentPercent":AccTreatyPaymentPercent.get(),
        "AccPayBodyCurTag":AccPayBodyCurTag.get()
        }

        template_document = Document(template_file_path)

        for variable_key, variable_value in variables.items():
            for paragraph in template_document.paragraphs:
                replace_text_in_paragraph(paragraph, variable_key, variable_value)

            for table in template_document.tables:
                for col in table.columns:
                    for cell in col.cells:
                        for paragraph in cell.paragraphs:
                            replace_text_in_paragraph(paragraph, variable_key, variable_value)

        template_document.save(output_file_path)
        msg_box = messagebox.showinfo("Успішно","Заміна у файлі успішно виконанна")
        if msg_box == 'ok':
            self.deptempusual.destroy()

         
    def on_closing(selt):
        if messagebox.askokcancel("Вийти?","Ви бажаєте закрити додаток?"):
            selt.deptempusual.destroy()
            
    
    def start(self):
           self.deptempusual.mainloop()           