from tkinter import *
from tkinter import ttk
from tkinter import messagebox
from docx import Document



def replace_text_in_paragraph(paragraph, key, value):
            if key in paragraph.text:
                inline = paragraph.runs
                for item in inline:
                    if key in item.text:
                        item.text = item.text.replace(key, value)  


class DepTem3x45():
    def __init__(self):
        self.deptep3x45 = Tk()
        self.deptep3x45.geometry("800x800")
        self.deptep3x45.title("KepHelper")
        self.deptep3x45.protocol("WM_DELETE_WINDOW", self.on_closing)

        global FileName
        global CurrentDate
        global ClientCode
        global TreatyCode
        global AccTreatyPlacing
        global AccPlacingCurTag
        global ClientName
        global StateCode
        global DepositIBAN
        global DepositAccCurrency
        global DepositAmountInWords
        global DepositAmount
        global RateInWords
        global Rate
        global Err
        global Ett
        global Rtt
        global Dtt
        global Epp
        global Eww
        global Rqq
        global Rbb
        global Eqq
        global Ebb
        global DateFrom
        global DateInto
        global DayCount
        global PercentReturnType
        global AccTreatyPaymentPercent
        global AccPayBodyCurTag
        global ClientOpenDate

        Label(self.deptep3x45,text='FileName').place(x=250,y=15)
        FileName = Entry(self.deptep3x45,width=50)
        FileName.place(x=250, y=35) 
        Label(self.deptep3x45, text="CurrentDate").place(x=50,y=60)
        CurrentDate = Entry(self.deptep3x45)
        CurrentDate.place(x=50, y=80)
        Label(self.deptep3x45, text="ClientCode").place(x=50,y=120)
        ClientCode = Entry(self.deptep3x45)
        ClientCode.place(x=50, y=140)
        Label(self.deptep3x45, text="TreatyCode").place(x=50,y=180)
        TreatyCode = Entry(self.deptep3x45)
        TreatyCode.place(x=50, y=200)
        Label(self.deptep3x45, text="AccTreatyPlacing").place(x=50,y=240)
        AccTreatyPlacing = Entry(self.deptep3x45)
        AccTreatyPlacing.place(x=50, y=260)
        Label(self.deptep3x45, text="AccPlacingCurTag").place(x=50,y=300)
        AccPlacingCurTag = Entry(self.deptep3x45)
        AccPlacingCurTag.place(x=50, y=320)
        Label(self.deptep3x45, text="ClientName").place(x=50,y=360)
        ClientName = Entry(self.deptep3x45)
        ClientName.place(x=50, y=380)
        Label(self.deptep3x45, text="StateCode").place(x=50,y=420)
        StateCode = Entry(self.deptep3x45)
        StateCode.place(x=50, y=440)
        Label(self.deptep3x45, text="DepositIBAN").place(x=50,y=480)
        DepositIBAN = Entry(self.deptep3x45)
        DepositIBAN.place(x=50, y=500)
        Label(self.deptep3x45, text="DepositAccCurrency").place(x=50,y=540)
        DepositAccCurrency = Entry(self.deptep3x45)
        DepositAccCurrency.place(x=50, y=560)
        Label(self.deptep3x45, text="DepositAmount").place(x=50,y=600)
        DepositAmount = Entry(self.deptep3x45)
        DepositAmount.place(x=50, y=620)
        Label(self.deptep3x45, text="DepositAmountInWords").place(x=250,y=60)
        DepositAmountInWords = Entry(self.deptep3x45)
        DepositAmountInWords.place(x=250, y=80)
        Label(self.deptep3x45, text="Rate1").place(x=250,y=120)
        Rate = Entry(self.deptep3x45)
        Rate.place(x=250, y=140)
        Label(self.deptep3x45, text="Rate1InWords").place(x=250,y=180)
        RateInWords = Entry(self.deptep3x45)
        RateInWords.place(x=250, y=200)
        Label(self.deptep3x45, text="DateFrom1").place(x=250,y=240)
        Err = Entry(self.deptep3x45)
        Err.place(x=250, y=260)
        Label(self.deptep3x45, text="DateInto1").place(x=250,y=300)
        Ett = Entry(self.deptep3x45)
        Ett.place(x=250, y=320)
        Label(self.deptep3x45, text="Rate2").place(x=250,y=360)
        Rtt = Entry(self.deptep3x45)
        Rtt.place(x=250, y=380)
        Label(self.deptep3x45, text="Rate2InWords").place(x=250,y=420)
        Dtt = Entry(self.deptep3x45)
        Dtt.place(x=250, y=440)
        Label(self.deptep3x45, text="DateFrom2").place(x=250,y=480)
        Epp = Entry(self.deptep3x45)
        Epp.place(x=250, y=500)
        Label(self.deptep3x45, text="DateInto2").place(x=250,y=540)
        Eww = Entry(self.deptep3x45)
        Eww.place(x=250, y=560)
        Label(self.deptep3x45, text="Rate3").place(x=250,y=600)
        Rqq = Entry(self.deptep3x45)
        Rqq.place(x=250, y=620)
        Label(self.deptep3x45, text="Rate3InWords").place(x=500,y=60)
        Rbb = Entry(self.deptep3x45)
        Rbb.place(x=500, y=80)
        Label(self.deptep3x45, text="DateFrom3").place(x=500,y=120)
        Eqq = Entry(self.deptep3x45)
        Eqq.place(x=500, y=140)
        Label(self.deptep3x45, text="DateInto3").place(x=500,y=180)
        Ebb = Entry(self.deptep3x45)
        Ebb.place(x=500, y=200)
        Label(self.deptep3x45, text="DateFrom").place(x=500,y=240)
        DateFrom = Entry(self.deptep3x45)
        DateFrom.place(x=500, y=260)
        Label(self.deptep3x45, text="DateInto").place(x=500,y=300)
        DateInto = Entry(self.deptep3x45)
        DateInto.place(x=500, y=320)
        Label(self.deptep3x45, text="DayCount").place(x=500,y=360)
        DayCount = Entry(self.deptep3x45)
        DayCount.place(x=500, y=380)
        Label(self.deptep3x45, text="PercentReturnType").place(x=500,y=420)
        PercentReturnType = Entry(self.deptep3x45)
        PercentReturnType.place(x=500, y=440)
        Label(self.deptep3x45, text="AccTreatyPaymentPercent").place(x=500,y=480)
        AccTreatyPaymentPercent = Entry(self.deptep3x45)
        AccTreatyPaymentPercent.place(x=500, y=500)
        Label(self.deptep3x45, text="AccPayBodyCurTag").place(x=500,y=540)
        AccPayBodyCurTag = Entry(self.deptep3x45)
        AccPayBodyCurTag.place(x=500, y=560)
        Label(self.deptep3x45, text="ClientOpenDate").place(x=500,y=600)
        ClientOpenDate = Entry(self.deptep3x45)
        ClientOpenDate.place(x=500, y=620)
        Button(self.deptep3x45,text="Replace Text",command=self.search,height=3,width=10).place(x= 500,y=730)



    def search(self):
        template_file_path = 'C:\Test\DepositTemplate3x45.docx'
        output_file_path = 'C:\Test\\' + FileName.get() + '.docx'

        variables = {
        "{":"",
        "}":"",
        "CurrentDate": CurrentDate.get(),
        "ClientCode": ClientCode.get(),
        "TreatyCode": TreatyCode.get(),
        "AccTreatyPlacing": AccTreatyPlacing.get(),
        "AccPlacingCurTag": AccPlacingCurTag.get(),
        "ClientName": ClientName.get(),
        "StateCode": StateCode.get(),
        "DepositIBAN": DepositIBAN.get(),
        "DepositAccCurrency": DepositAccCurrency.get(),
        "DepositAmountInWords": DepositAmountInWords.get(),
        "DepositAmount": DepositAmount.get(),
        "RateInWords": RateInWords.get(),
        "Rate": Rate.get(),
        "Err": Err.get(),
        "Ett": Ett.get(),
        "Rtt": Rtt.get(),
        "Dtt": Dtt.get(),
        "Epp": Epp.get(),
        "Eww": Eww.get(),
        "Rqq": Rqq.get(),
        "Rbb":Rbb.get(),
        "Eqq":Eqq.get(),
        "Ebb":Ebb.get(),
        "DateFrom":DateFrom.get(),
        "DateInto":DateInto.get(),
        "DayCount":DayCount.get(),
        "PercentReturnType": PercentReturnType.get(),
        "AccTreatyPaymentPercent": AccTreatyPaymentPercent.get(),
        "AccPayBodyCurTag": AccPayBodyCurTag.get(),
        "ClientOpenDate":ClientOpenDate.get()
        }
        template_document = Document(template_file_path)

        for variable_key, variable_value in variables.items():
            for paragraph in template_document.paragraphs:
                replace_text_in_paragraph(paragraph, variable_key, variable_value)

            for table in template_document.tables:
                for col in table.columns:
                    for cell in col.cells:
                        for paragraph in cell.paragraphs:
                            replace_text_in_paragraph(paragraph, variable_key, variable_value)

        template_document.save(output_file_path)
        msg_box = messagebox.showinfo("Успішно","Заміна у файлі успішно виконанна")
        if msg_box == 'ok':
            self.deptep3x45.destroy()
    def on_closing(selt):
        if messagebox.askokcancel("Вийти?","Ви бажаєте закрити додаток?"):
            selt.deptep3x45.destroy()
            
            
    
    def start(self):
           self.deptep3x45.mainloop()
