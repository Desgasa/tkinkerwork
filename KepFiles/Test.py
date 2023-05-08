from tkinter import *
from tkinter import ttk
from tkinter import messagebox
from docx import Document

from database.configdatabase import *





def replace_text_in_paragraph(paragraph, key, value):
            if key in paragraph.text:
                inline = paragraph.runs
                for item in inline:
                    if key in item.text:
                        item.text = item.text.replace(key, value)  

class test():
    def __init__(self):
        global Treatyid
        global Filename
        self.test = Tk()
        self.test.geometry("300x250")
        self.test.title("KepHelper")
        self.test.protocol("WM_DELETE_WINDOW", self.on_closing)
        Label(self.test, text = 'FileName').place(x=87,y=40)
        Filename = Entry(self.test, width = 30)
        Filename.place(x=87, y=65)
        Label(self.test,text='Treatyid').place(x=87,y=100)
        Treatyid = Entry(self.test,width=30)
        Treatyid.place(x=87, y=125)
        replacebutton = ttk.Button(self.test,text="Replace Text",command=self.search)
        replacebutton.pack(
            
            expand= True,
            ipadx = 2, 
            ipady = 7,
            anchor = "s",
            pady = 30
            )
        

        

    def search(self):
        mycursor = dbconn.cursor()
        args = [Treatyid.get(), 0]
        result = mycursor.callproc('p2', args)
        datas = [Treatyid.get(),0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0]
        tester = mycursor.callproc('test', datas)




        template_file_path = 'C:\Test\\' + result[1] + '.docx'
        output_file_path = 'C:\Test\\' +  Filename.get()  + '.docx'
        variables = {
        "{":"",
        "}":"",
        "CurrentDate": tester[22],
        "ClientCode": tester[1],
        "TreatyCode": tester[21],
        "ClientName": tester[2],
        "StateCode": tester[4],
        "ClientAddress":tester[5],
        "AccTreatyPlacing":tester[6],
        "AccPlacingCurTag":tester[7],
        "DepositIBAN":tester[8],
        "DepositAccCurrency":tester[9],
        "DepositAmountInWords":tester[11],
        "DepositAmount":tester[10],
        "RateInWords":tester[13],
        "Rate": tester[12],
        "DateFrom":tester[14],
        "DateInto":tester[15],
        "DayCount":tester[16],
        "AutolongTreaty":tester[17],
        "PercentReturnType":tester[18],
        "AccTreatyPaymentPercent":tester[19],
        "AccPayBodyCurTag":tester[20]
                                            }



        template_document = Document(template_file_path)

        for variable_key, variable_value in variables.items():
            for paragraph in template_document.paragraphs:
                replace_text_in_paragraph(paragraph, variable_key, variable_value)

            for table in template_document.tables:
                for col in table.columns:
                    for cell in col.cells:
                        for paragraph in cell.paragraphs:
                            replace_text_in_paragraph(paragraph, variable_key, variable_value)

        template_document.save(output_file_path)



    def on_closing(self):
        if messagebox.askokcancel("Вийти?","Ви бажаєте закрити додаток?"):
            self.test2.destroy()
            
    
    def start(self):
           self.test2.mainloop()
           

         