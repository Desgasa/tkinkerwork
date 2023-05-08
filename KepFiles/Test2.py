from tkinter import *
from tkinter import ttk
from tkinter import messagebox
from docx import Document

from database.configdatabaseserver import *

def replace_text_in_paragraph(paragraph, key, value):
            if key in paragraph.text:
                inline = paragraph.runs
                for item in inline:
                    if key in item.text:
                        item.text = item.text.replace(key, value)  


class test2():
    def __init__(self):
        global Treatyid
        global Filename
        self.test2 = Tk()
        self.test2.geometry("300x250")
        self.test2.title("KepHelper")
        self.test2.protocol("WM_DELETE_WINDOW", self.on_closing)
        Label(self.test2, text = 'FileName').place(x=87,y=40)
        Filename = Entry(self.test2, width = 30)
        Filename.place(x=87, y=65)
        Label(self.test2,text='Treatyid').place(x=87,y=100)
        Treatyid = Entry(self.test2,width=30)
        Treatyid.place(x=87, y=125)
        replacebutton = ttk.Button(self.test2,text="Replace Text",command=self.search)
        replacebutton.pack(
            
            expand= True,
            ipadx = 2, 
            ipady = 7,
            anchor = "s",
            pady = 30
            )
        

    def search(self):


        mycursor = connection.cursor()
        result = mycursor.execute(f"EXEC getTemplate @test = {Treatyid.get()}, @template = OUTPUT;").fetchone()
        connection.commit()

        if result[0] == "DepositTemplateUsual":

            test = mycursor.execute(f'exec getDatas @test = {Treatyid.get()}, @datas1 = OUTPUT, @datas2 = OUTPUT;').fetchone()


            template_file_path = 'C:\Test\\' + result[0] + '.docx'
            output_file_path = 'C:\Test\\' +  Filename.get()  + '.docx'
            variables = {
            "{":"",
            "}":"",
            "CurrentDate": test[0],
            "ClientCode": test[1]}
            template_document = Document(template_file_path)

            for variable_key, variable_value in variables.items():
                for paragraph in template_document.paragraphs:
                    replace_text_in_paragraph(paragraph, variable_key, variable_value)

                for table in template_document.tables:
                    for col in table.columns:
                        for cell in col.cells:
                            for paragraph in cell.paragraphs:
                                replace_text_in_paragraph(paragraph, variable_key, variable_value)
            template_document.save(output_file_path)
        elif result[0] == "DepositTemplateEarlyDessolution":
            test = mycursor.execute(f'exec getDatas @test = {Treatyid.get()}, @datas1 = OUTPUT, @datas2 = OUTPUT;').fetchone()


            template_file_path = 'C:\Test\\' + result[0] + '.docx'
            output_file_path = 'C:\Test\\' +  Filename.get()  + '.docx'
            variables = {
            "{":"",
            "}":"",
            "CurrentDate": test[0],
            "ClientCode": test[1]}
            template_document = Document(template_file_path)

            for variable_key, variable_value in variables.items():
                for paragraph in template_document.paragraphs:
                    replace_text_in_paragraph(paragraph, variable_key, variable_value)

                for table in template_document.tables:
                    for col in table.columns:
                        for cell in col.cells:
                            for paragraph in cell.paragraphs:
                                replace_text_in_paragraph(paragraph, variable_key, variable_value)
            template_document.save(output_file_path)


        elif result[0] == "DepositTemplate3x20":
            test = mycursor.execute(f'exec getDatas @test = {Treatyid.get()}, @datas1 = OUTPUT, @datas2 = OUTPUT;').fetchone()


            template_file_path = 'C:\Test\\' + result[0] + '.docx'
            output_file_path = 'C:\Test\\' +  Filename.get()  + '.docx'
            variables = {
            "{":"",
            "}":"",
            "CurrentDate": test[0],
            "ClientCode": test[1]}
            template_document = Document(template_file_path)

            for variable_key, variable_value in variables.items():
                for paragraph in template_document.paragraphs:
                    replace_text_in_paragraph(paragraph, variable_key, variable_value)

                for table in template_document.tables:
                    for col in table.columns:
                        for cell in col.cells:
                            for paragraph in cell.paragraphs:
                                replace_text_in_paragraph(paragraph, variable_key, variable_value)
            template_document.save(output_file_path)     
             
        elif result[0] == "DepositTemplate3x45":
            test = mycursor.execute(f'exec getDatas @test = {Treatyid.get()}, @datas1 = OUTPUT, @datas2 = OUTPUT;').fetchone()


            template_file_path = 'C:\Test\\' + result[0] + '.docx'
            output_file_path = 'C:\Test\\' +  Filename.get()  + '.docx'
            variables = {
            "{":"",
            "}":"",
            "CurrentDate": test[0],
            "ClientCode": test[1]}
            template_document = Document(template_file_path)

            for variable_key, variable_value in variables.items():
                for paragraph in template_document.paragraphs:
                    replace_text_in_paragraph(paragraph, variable_key, variable_value)

                for table in template_document.tables:
                    for col in table.columns:
                        for cell in col.cells:
                            for paragraph in cell.paragraphs:
                                replace_text_in_paragraph(paragraph, variable_key, variable_value)
            template_document.save(output_file_path)     
             


    def on_closing(self):
        if messagebox.askokcancel("Вийти?","Ви бажаєте закрити додаток?"):
            self.test2.destroy()
            
    
    def start(self):
           self.test2.mainloop()
           