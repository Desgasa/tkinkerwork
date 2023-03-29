from tkinter import *
from tkinter import ttk
from tkinter import messagebox
from docx import Document





def replace_text_in_paragraph(paragraph, key, value):
            if key in paragraph.text:
                inline = paragraph.runs
                for item in inline:
                    if key in item.text:
                        item.text = item.text.replace(key, value)  
class DepTED():
    def __init__(self):
        self.deptempearldess = Tk()
        self.deptempearldess.geometry("800x800")
        self.deptempearldess.title("KepHelper")
        self.deptempearldess.protocol("WM_DELETE_WINDOW", self.on_closing)

        global FileName
        global CurrentDate
        global ClientCode
        global TreatyCode
        global ClientName
        global StateCode
        global ClientAddress
        global AccTreatyPlacing
        global AccPlacingCurTag
        global DepositIBAN
        global DepositAccCurrency
        global DepositAmount
        global DepositAmountInWords
        global Rate
        global RateInWords
        global DateFrom
        global DateInto
        global DayCount
        global AutolongTreaty
        global PercentReturnType
        global AccTreatyPaymentPercent
        global AccPayBodyCurTag

        Label(self.deptempearldess,text='FileName').place(x=250,y=15)
        FileName = Entry(self.deptempearldess,width=50)
        FileName.place(x=250, y=35)
        Label(self.deptempearldess, text="CurrentDate").place(x=50,y=60)
        CurrentDate = Entry(self.deptempearldess)
        CurrentDate.place(x=50, y=80)
        Label(self.deptempearldess, text="ClientCode").place(x=50,y=120)
        ClientCode = Entry(self.deptempearldess)
        ClientCode.place(x=50, y=140)
        Label(self.deptempearldess, text="TreatyCode").place(x=50,y=180)
        TreatyCode = Entry(self.deptempearldess)
        TreatyCode.place(x=50, y=200)
        Label(self.deptempearldess, text="ClientName").place(x=50,y=240)
        ClientName = Entry(self.deptempearldess)
        ClientName.place(x=50, y=260)
        Label(self.deptempearldess, text="StateCode").place(x=50,y=300)
        StateCode = Entry(self.deptempearldess)
        StateCode.place(x=50, y=320)
        Label(self.deptempearldess, text="ClientAddress").place(x=50,y=360)
        ClientAddress = Entry(self.deptempearldess)
        ClientAddress.place(x=50, y=380)
        Label(self.deptempearldess, text="AccTreatyPlacing").place(x=50,y=420)
        AccTreatyPlacing = Entry(self.deptempearldess)
        AccTreatyPlacing.place(x=50, y=440)
        Label(self.deptempearldess, text="AccPlacingCurTag").place(x=50,y=480)
        AccPlacingCurTag = Entry(self.deptempearldess)
        AccPlacingCurTag.place(x=50, y=500)
        Label(self.deptempearldess, text="DepositIBAN").place(x=50,y=540)
        DepositIBAN = Entry(self.deptempearldess)
        DepositIBAN.place(x=50, y=560)
        Label(self.deptempearldess, text="DepositAccCurrency").place(x=50,y=600)
        DepositAccCurrency = Entry(self.deptempearldess)
        DepositAccCurrency.place(x=50, y=620)
        Label(self.deptempearldess, text="DepositAmount").place(x=250,y=60)
        DepositAmount = Entry(self.deptempearldess)
        DepositAmount.place(x=250, y=80)
        Label(self.deptempearldess, text="DepositAmountInWords").place(x=250,y=120)
        DepositAmountInWords = Entry(self.deptempearldess)
        DepositAmountInWords.place(x=250, y=140)
        Label(self.deptempearldess, text="Rate").place(x=250,y=180)
        Rate = Entry(self.deptempearldess)
        Rate.place(x=250, y=200)
        Label(self.deptempearldess, text="RateInWords").place(x=250,y=240)
        RateInWords = Entry(self.deptempearldess)
        RateInWords.place(x=250, y=260)
        Label(self.deptempearldess, text="DateFrom").place(x=250,y=300)
        DateFrom = Entry(self.deptempearldess)
        DateFrom.place(x=250, y=320)
        Label(self.deptempearldess, text="DateInto").place(x=250,y=360)
        DateInto = Entry(self.deptempearldess)
        DateInto.place(x=250, y=380)
        Label(self.deptempearldess, text="DayCount").place(x=250,y=420)
        DayCount = Entry(self.deptempearldess)
        DayCount.place(x=250, y=440)
        Label(self.deptempearldess, text="AutolongTreaty").place(x=250,y=480)
        AutolongTreaty = Entry(self.deptempearldess)
        AutolongTreaty.place(x=250, y=500)
        Label(self.deptempearldess, text="PercentReturnType").place(x=250,y=540)
        PercentReturnType = Entry(self.deptempearldess)
        PercentReturnType.place(x=250, y=560)
        Label(self.deptempearldess, text="AccTreatyPaymentPercent").place(x=250,y=600)
        AccTreatyPaymentPercent = Entry(self.deptempearldess)
        AccTreatyPaymentPercent.place(x=250, y=620)
        Label(self.deptempearldess, text="AccPayBodyCurTag").place(x=500,y=60)
        AccPayBodyCurTag = Entry(self.deptempearldess)
        AccPayBodyCurTag.place(x=500, y=80)
        Button(self.deptempearldess,text="Replace Text",command = self.search,height=3,width=10).place(x= 500,y=730)
    
    def search(self):
        template_file_path = 'C:\Test\DepositTemplateUsual.docx'
        output_file_path = 'C:\Test\\' + FileName.get() + '.docx'

        variables = {
        "{":"",
        "}":"",
        "CurrentDate": CurrentDate.get(),
        "ClientCode": ClientCode.get(),
        "TreatyCode": TreatyCode.get(),
        "ClientName": ClientName.get(),
        "StateCode": StateCode.get(),
        "ClientAddress": ClientAddress.get(),
        "AccTreatyPlacing": AccTreatyPlacing.get(),
        "AccPlacingCurTag":AccPlacingCurTag.get(),
        "DepositIBAN":DepositIBAN.get(),
        "DepositAccCurrency":DepositAccCurrency.get(),
        "DepositAmountInWords":DepositAmountInWords.get(),
        "DepositAmount":DepositAmount.get(),
        "RateInWords":RateInWords.get(),
        "Rate":Rate.get(),
        "DateFrom":DateFrom.get(),
        "DateInto":DateInto.get(),
        "DayCount":DayCount.get(),
        "AutolongTreaty":AutolongTreaty.get(),
        "PercentReturnType":PercentReturnType.get(),
        "AccTreatyPaymentPercent":AccTreatyPaymentPercent.get(),
        "AccPayBodyCurTag":AccPayBodyCurTag.get()
        }

        template_document = Document(template_file_path)

        for variable_key, variable_value in variables.items():
            for paragraph in template_document.paragraphs:
                replace_text_in_paragraph(paragraph, variable_key, variable_value)

            for table in template_document.tables:
                for col in table.columns:
                    for cell in col.cells:
                        for paragraph in cell.paragraphs:
                            replace_text_in_paragraph(paragraph, variable_key, variable_value)

        template_document.save(output_file_path)
        msg_box = messagebox.showinfo("Успішно","Заміна у файлі успішно виконанна")
        if msg_box == 'ok':
            self.deptempusual.destroy()
    def on_closing(selt):
        if messagebox.askokcancel("Вийти?","Ви бажаєте закрити додаток?"):
            selt.deptempearldess.destroy()
            
    
    def start(self):
           self.deptempearldess.mainloop()           