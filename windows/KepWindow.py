from tkinter import *
from tkinter import ttk
from tkinter import messagebox
from docx import Document

from database.configdatabaseserver import *


def replace_text_in_paragraph(paragraph, key, value):
            if key in paragraph.text:
                inline = paragraph.runs
                for item in inline:
                    if key in item.text:
                        item.text = item.text.replace(key, value)  



class root():
    def __init__(self):
        global Treatyid
        global Filename
        self.root = Tk()
        self.root.geometry("300x250")
        self.root.title("KepHelper")
        self.root.protocol("WM_DELETE_WINDOW", self.on_closing)
        Label(self.root, text = 'FileName').place(x=87,y=40)
        Filename = Entry(self.root, width = 30)
        Filename.place(x=87, y=65)
        Label(self.root,text='Treatyid').place(x=87,y=100)
        Treatyid = Entry(self.root,width=30)
        Treatyid.place(x=87, y=125)
        replacebutton = ttk.Button(self.root,text="Replace Text",command= self.search)
        replacebutton.pack(
            
            expand= True,
            ipadx = 2, 
            ipady = 7,
            anchor = "s",
            pady = 30
            )
        

    def search(self):


        mycursor = connection.cursor()
        template = mycursor.execute(f"EXEC getTemplate @test = {Treatyid.get()}, @template = OUTPUT;").fetchone()
        connection.commit()

        if template[0] == "DepositTemplateUsual":

            test = mycursor.execute(f'exec getDatasTempUsual @treatyid = {Treatyid.get()}, @CurrentDate = OUTPUT, @ClientCode = OUTPUT,@TreatyCode = OUTPUT, @ClientName = OUTPUT, @StateCode = OUTPUT, @ClientAddress = OUTPUT, @AccTreatyPlacing = OUTPUT,@AccPlacingCurTag = OUTPUT,@DepositIBAN = OUTPUT, @DepositAccCurrency = OUTPUT, @DepositAmount = OUTPUT, @DepositAmountInWords = OUTPUT, @Rate = OUTPUT, @RateInWords = OUTPUT, @DateFrom = OUTPUT, @DateInto = OUTPUT, @DayCount = OUTPUT, @AutolongTreaty = OUTPUT, @PercentReturnType = OUTPUT,@AccTreatyPaymentPercent = OUTPUT, @AccPayBodyCurTag = OUTPUT;').fetchone()



            template_file_path = 'C:\Test\\' + template[0] + '.docx'
            output_file_path = 'C:\Test\\' +  Filename.get()  + '.docx'
            variables = {
            "{":"",
            "}":"",
            "CurrentDate": test[0],
            "ClientCode": test[1],
            'TreatyCode': test[2],
            "ClientName": test[3],
            'StateCode': test[4],
            'ClientAddress': test[5],
            'AccTreatyPlacing':test[6],
            'AccPlacingCurTag':test[7],
            'DepositIBAN': test[8],
            'DepositAccCurrency':test[9],
            'DepositAmountInWords':test[11],
            'DepositAmount': test[10],
            "RateInWords":test[13],
            "Rate": test[12],
            "DateFrom":test[14],
            "DateInto":test[15],
            "DayCount":test[16],
            "AutolongTreaty":test[17],
            "PercentReturnType":test[18],
            "AccTreatyPaymentPercent":test[19],
            "AccPayBodyCurTag":test[20]
            }
            template_document = Document(template_file_path)

            for variable_key, variable_value in variables.items():
                for paragraph in template_document.paragraphs:
                    replace_text_in_paragraph(paragraph, variable_key, variable_value)

                for table in template_document.tables:
                    for col in table.columns:
                        for cell in col.cells:
                            for paragraph in cell.paragraphs:
                                replace_text_in_paragraph(paragraph, variable_key, variable_value)
            template_document.save(output_file_path)

            
        elif template[0] == "DepositTemplateEarlyDessolution":


            test = mycursor.execute(f'exec getDatasTempUsual @treatyid = {Treatyid.get()}, @CurrentDate = OUTPUT, @ClientCode = OUTPUT,@TreatyCode = OUTPUT, @ClientName = OUTPUT, @StateCode = OUTPUT, @ClientAddress = OUTPUT, @AccTreatyPlacing = OUTPUT,@AccPlacingCurTag = OUTPUT,@DepositIBAN = OUTPUT, @DepositAccCurrency = OUTPUT, @DepositAmount = OUTPUT, @DepositAmountInWords = OUTPUT, @Rate = OUTPUT, @RateInWords = OUTPUT, @DateFrom = OUTPUT, @DateInto = OUTPUT, @DayCount = OUTPUT, @AutolongTreaty = OUTPUT, @PercentReturnType = OUTPUT,@AccTreatyPaymentPercent = OUTPUT, @AccPayBodyCurTag = OUTPUT;').fetchone()



            template_file_path = 'C:\Test\\' + template[0] + '.docx'
            output_file_path = 'C:\Test\\' +  Filename.get()  + '.docx'
            variables = {
             "{":"",
            "}":"",
            "CurrentDate": test[0],
            "ClientCode": test[1],
            'TreatyCode': test[2],
            "ClientName": test[3],
            'StateCode': test[4],
            'ClientAddress': test[5],
            'AccTreatyPlacing':test[6],
            'AccPlacingCurTag':test[7],
            'DepositIBAN': test[8],
            'DepositAccCurrency':test[9],
            'DepositAmountInWords':test[11],
            'DepositAmount': test[10],
            "RateInWords":test[13],
            "Rate": test[12],
            "DateFrom":test[14],
            "DateInto":test[15],
            "DayCount":test[16],
            "AutolongTreaty":test[17],
            "PercentReturnType":test[18],
            "AccTreatyPaymentPercent":test[19],
            "AccPayBodyCurTag":test[20]
            }
            template_document = Document(template_file_path)

            for variable_key, variable_value in variables.items():
                for paragraph in template_document.paragraphs:
                    replace_text_in_paragraph(paragraph, variable_key, variable_value)

                for table in template_document.tables:
                    for col in table.columns:
                        for cell in col.cells:
                            for paragraph in cell.paragraphs:
                                replace_text_in_paragraph(paragraph, variable_key, variable_value)
            template_document.save(output_file_path)


        elif template[0] == "DepositTemplate3x20":
            test = mycursor.execute(f'exec getDataTempl3x @treatyid = {Treatyid.get()}, @CurrentDate = output, @ClientCode = OUTPUT, @TreatyCode = OUTPUT, @AccTreatyPlacing = output, @AccPlacingCurTag = output,@ClientName = output,@StateCode = output,@DepositIBAN = output,@DepositAccCurrency = output,@DepositAmountInWords = output,@DepositAmount = output,@Rate1InWords = output, @Rate1 = output,@DateFrom1 = output,@DateInto1 = output,@Rate2InWords = output,@Rate2 = output, @DateFrom2 = output, @DateInto2 = output, @Rate3InWords = output, @Rate3 = output,@DateFrom3 = output, @DateInto3 = output,@DateFrom = output,@DateInto = output, @DayCount = output, @PercentReturnType = output, @AccTreatyPaymentPercent = output, @AccPayBodyCurTag = output, @ClientOpenDate = output;').fetchone()


            template_file_path = 'C:\Test\\' + template[0] + '.docx'
            output_file_path = 'C:\Test\\' +  Filename.get()  + '.docx'
            variables = {
            "{":"",
            "}":"",
            "CurrentDate": test[0],
            "ClientCode": test[1],
            "TreatyCode": test[2],
            'AccTreatyPlacing': test[3],
            'AccPlacingCurTag': test[4],
            'ClientName': test[5],
            'StateCode':test[6],
            'DepositIBAN':test[7],
            'DepositAccCurrency':test[8],
            "DepositAmountInWords":test[9],
            "DepositAmount":test[10],
            "RateInWords": test[11],
            'Rate': test[12],
            'Err': test[13],
            'Ett': test[14],
            "Dtt": test[15],
            'Rtt': test[16],
            "Epp": test[17],
            "Eww":test[18],
            "Rbb": test[19],
            "Rqq": test[20],
            "Eqq": test[21],
            "Ebb":test[22],
            'DateFrom':test[23],
            'DateInto': test[24],
            'DayCount': test[25],
            "PercentReturnType": test[26],
            'AccTreatyPaymentPercent': test[27],
            'AccPayBodyCurTag': test[28],
            'ClientOpenDate': test[29]}
            

            template_document = Document(template_file_path)

            for variable_key, variable_value in variables.items():
                for paragraph in template_document.paragraphs:
                    replace_text_in_paragraph(paragraph, variable_key, variable_value)

                for table in template_document.tables:
                    for col in table.columns:
                        for cell in col.cells:
                            for paragraph in cell.paragraphs:
                                replace_text_in_paragraph(paragraph, variable_key, variable_value)
            template_document.save(output_file_path)     
             
        elif template[0] == "DepositTemplate3x45":
            test = mycursor.execute(f'exec getDataTempl3x @treatyid = {Treatyid.get()}, @CurrentDate = output, @ClientCode = OUTPUT, @TreatyCode = OUTPUT, @AccTreatyPlacing = output, @AccPlacingCurTag = output,@ClientName = output,@StateCode = output,@DepositIBAN = output,@DepositAccCurrency = output,@DepositAmountInWords = output,@DepositAmount = output,@Rate1InWords = output, @Rate1 = output,@DateFrom1 = output,@DateInto1 = output,@Rate2InWords = output,@Rate2 = output, @DateFrom2 = output, @DateInto2 = output, @Rate3InWords = output, @Rate3 = output,@DateFrom3 = output, @DateInto3 = output,@DateFrom = output,@DateInto = output, @DayCount = output, @PercentReturnType = output, @AccTreatyPaymentPercent = output, @AccPayBodyCurTag = output, @ClientOpenDate = output;').fetchone()


            template_file_path = 'C:\Test\\' + template[0] + '.docx'
            output_file_path = 'C:\Test\\' +  Filename.get()  + '.docx'
            variables = {
            "{":"",
            "}":"",
            "CurrentDate": test[0],
            "ClientCode": test[1],
            "TreatyCode": test[2],
            'AccTreatyPlacing': test[3],
            'AccPlacingCurTag': test[4],
            'ClientName': test[5],
            'StateCode':test[6],
            'DepositIBAN':test[7],
            'DepositAccCurrency':test[8],
            "DepositAmountInWords":test[9],
            "DepositAmount":test[10],
            "RateInWords": test[11],
            'Rate': test[12],
            'Err': test[13],
            'Ett': test[14],
            "Dtt": test[15],
            'Rtt': test[16],
            "Epp": test[17],
            "Eww":test[18],
            "Rbb": test[19],
            "Rqq": test[20],
            "Eqq": test[21],
            "Ebb":test[22],
            'DateFrom':test[23],
            'DateInto': test[24],
            'DayCount': test[25],
            "PercentReturnType": test[26],
            'AccTreatyPaymentPercent': test[27],
            'AccPayBodyCurTag': test[28],
            'ClientOpenDate': test[29]}
            template_document = Document(template_file_path)

            for variable_key, variable_value in variables.items():
                for paragraph in template_document.paragraphs:
                    replace_text_in_paragraph(paragraph, variable_key, variable_value)

                for table in template_document.tables:
                    for col in table.columns:
                        for cell in col.cells:
                            for paragraph in cell.paragraphs:
                                replace_text_in_paragraph(paragraph, variable_key, variable_value)
            template_document.save(output_file_path)     
             

    def on_closing(self):
        if messagebox.askokcancel("Вийти?","Ви бажаєте закрити додаток?"):
            self.root.destroy()
            
    
    def start(self):
           self.root.mainloop()
               